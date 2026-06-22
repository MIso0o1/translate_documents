import docx
import asyncio
import httpx
from translate_text import translate_text_with_deepl_async

def apply_formatting(run, formatting):
    if formatting:
        run.bold = True

async def translate_run_async(run, target_lang, client):
    original_text = run.text
    if original_text and original_text.strip():
        translated_text = await translate_text_with_deepl_async(original_text, target_lang, client)
        if translated_text:
            run.clear()
            run.text = translated_text
            apply_formatting(run, run.bold)

async def translate_paragraph_async(paragraph, target_lang, client):
    tasks = [translate_run_async(run, target_lang, client) for run in paragraph.runs]
    if tasks:
        await asyncio.gather(*tasks)

async def translate_tables_in_docx_async(doc, target_lang, client):
    tasks = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    tasks.append(translate_paragraph_async(paragraph, target_lang, client))
    if tasks:
        await asyncio.gather(*tasks)

def clean_toc_entry(entry):
    entry = entry.split('\t', 1)[-1].strip()
    entry = entry.split('\\o', 1)[0].strip()
    return entry

async def translate_docx_async(docx_file_path, target_lang):
    doc = docx.Document(docx_file_path)
    toc_entries = []

    # Identify TOC entries
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("toc"):
            toc_entries.append(clean_toc_entry(paragraph.text))

    async with httpx.AsyncClient() as client:
        # Translate TOC entries
        if toc_entries:
            toc_tasks = [translate_text_with_deepl_async(entry, target_lang, client) for entry in toc_entries]
            translated_toc_entries = await asyncio.gather(*toc_tasks)
            
            # Map back to paragraphs
            toc_idx = 0
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("toc") and toc_idx < len(translated_toc_entries):
                    if translated_toc_entries[toc_idx]:
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.text = translated_toc_entries[toc_idx]
                    toc_idx += 1

        # Translate normal paragraphs
        para_tasks = []
        for paragraph in doc.paragraphs:
            if not paragraph.style.name.startswith("toc"):
                para_tasks.append(translate_paragraph_async(paragraph, target_lang, client))
        
        if para_tasks:
            await asyncio.gather(*para_tasks)

        # Translate tables
        await translate_tables_in_docx_async(doc, target_lang, client)

    return doc

# Keep the original function for compatibility if needed, or wrap the async one
def translate_docx(docx_file_path, target_lang):
    return asyncio.run(translate_docx_async(docx_file_path, target_lang))
